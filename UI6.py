import tkinter as tk
from tkinter import messagebox, scrolledtext, filedialog
import pandas as pd
from sklearn.preprocessing import OneHotEncoder
from sklearn.ensemble import RandomForestRegressor
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

# -----------------------------
# 1️⃣ - Partie Modèle Machine Learning
# -----------------------------
data = pd.DataFrame({
    'Jour': ['Lundi','Lundi','Mardi','Mercredi','Jeudi','Vendredi'],
    'Heure': ['Matin','Soir','Matin','Midi','Soir','Matin'],
    'Meteo': ['soleil','pluie','soleil','neige','pluie','soleil'],
    'Promo': ['Oui','Non','Non','Oui','Oui','Non'],
    'Evenement': ['Non','Oui','Non','Non','Oui','Non'],
    'Clients': [120,180,140,100,200,160]
})

encoder = OneHotEncoder(sparse_output=False)
X_cat = encoder.fit_transform(data[['Jour','Heure','Meteo','Promo','Evenement']])
X_cat_df = pd.DataFrame(X_cat, columns=encoder.get_feature_names_out())
X = X_cat_df
y = data['Clients']

model = RandomForestRegressor(n_estimators=100, random_state=42)
model.fit(X, y)

# -----------------------------
# 2️⃣ - Interface Graphique Tkinter
# -----------------------------
root = tk.Tk()
root.title("Prédiction du nombre de clients")
root.geometry("700x700")
root.resizable(False, False)
root.configure(bg="#f8f9fa")

LABEL_FONT = ("Helvetica", 12)
TITLE_FONT = ("Helvetica", 18, "bold")
BUTTON_FONT = ("Helvetica", 12, "bold")
BG_COLOR = "#f8f9fa"
BTN_COLOR = "#0078D7"
BTN_TEXT_COLOR = "#ffffff"

tk.Label(root, text=" Prédiction du Nombre de Clients", font=TITLE_FONT, bg=BG_COLOR, fg="#333333").pack(pady=15)

frame_inputs = tk.Frame(root, bg=BG_COLOR)
frame_inputs.pack(pady=10)

jours = ['Lundi','Mardi','Mercredi','Jeudi','Vendredi']
heures = ['Matin','Midi','Soir']
meteo_options = ['soleil','pluie','neige']
promo_options = ['Oui','Non']
evenement_options = ['Oui','Non']

jour_var = tk.StringVar(value=jours[0])
heure_var = tk.StringVar(value=heures[0])
meteo_var = tk.StringVar(value=meteo_options[0])
promo_var = tk.StringVar(value=promo_options[0])
event_var = tk.StringVar(value=evenement_options[0])

options = [
    ("Jour", jour_var, jours),
    ("Heure", heure_var, heures),
    ("Météo", meteo_var, meteo_options),
    ("Promotion", promo_var, promo_options),
    ("Événement", event_var, evenement_options)
]

for i, (label, var, values) in enumerate(options):
    tk.Label(frame_inputs, text=label + " :", font=LABEL_FONT, bg=BG_COLOR).grid(row=i, column=0, padx=10, pady=8, sticky="w")
    tk.OptionMenu(frame_inputs, var, *values).grid(row=i, column=1, padx=10, pady=8, sticky="ew")

# -----------------------------
# 3️⃣ - Fonction de Prédiction
# -----------------------------
def predire():
    try:
        nouveau_jour = pd.DataFrame({
            'Jour': [jour_var.get()],
            'Heure': [heure_var.get()],
            'Meteo': [meteo_var.get()],
            'Promo': [promo_var.get()],
            'Evenement': [event_var.get()]
        })

        nouveau_jour_cat = encoder.transform(nouveau_jour)
        nouveau_jour_df = pd.DataFrame(nouveau_jour_cat, columns=encoder.get_feature_names_out())

        resultat = model.predict(nouveau_jour_df)
        prediction = int(resultat[0])

        label_result.config(text=f"✅ Clients prévus : {prediction}", fg="green")
        historique_text.insert(tk.END, f"{jour_var.get()} - {heure_var.get()} - {meteo_var.get()} - "
                                       f"Promo:{promo_var.get()} - Événement:{event_var.get()} ➜ {prediction} clients\n")
        historique_text.see(tk.END)
    except Exception as e:
        messagebox.showerror("Erreur", f"Une erreur est survenue : {e}")

# -----------------------------
# 4️⃣ - Réinitialiser
# -----------------------------
def reset():
    jour_var.set(jours[0])
    heure_var.set(heures[0])
    meteo_var.set(meteo_options[0])
    promo_var.set(promo_options[0])
    event_var.set(evenement_options[0])
    label_result.config(text="")
    historique_text.delete("1.0", tk.END)

# -----------------------------
# 5️⃣ - Fonctions Import / Export
# -----------------------------
def export_csv():
    try:
        data = historique_text.get("1.0", tk.END).strip().split("\n")
        if not data or data == ['']:
            messagebox.showwarning("Attention", "Aucune donnée à exporter.")
            return
        df = pd.DataFrame(data, columns=["Historique"])
        filepath = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("Fichier CSV", "*.csv")])
        if filepath:
            df.to_csv(filepath, index=False)
            messagebox.showinfo("Succès", "Export CSV réussi ✅")
    except Exception as e:
        messagebox.showerror("Erreur", str(e))

def import_csv():
    try:
        filepath = filedialog.askopenfilename(filetypes=[("Fichier CSV", "*.csv")])
        if filepath:
            df = pd.read_csv(filepath)
            historique_text.delete("1.0", tk.END)
            for row in df["Historique"]:
                historique_text.insert(tk.END, row + "\n")
            messagebox.showinfo("Succès", "Import CSV réussi ✅")
    except Exception as e:
        messagebox.showerror("Erreur", str(e))

def export_pdf():
    try:
        data = historique_text.get("1.0", tk.END).strip()
        if not data:
            messagebox.showwarning("Attention", "Aucune donnée à exporter.")
            return
        filepath = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("Fichier PDF", "*.pdf")])
        if filepath:
            doc = SimpleDocTemplate(filepath, pagesize=A4)
            styles = getSampleStyleSheet()
            story = [Paragraph("Rapport des Prédictions", styles['Title']), Spacer(1, 12)]
            for line in data.split("\n"):
                story.append(Paragraph(line, styles['Normal']))
            doc.build(story)
            messagebox.showinfo("Succès", "Export PDF réussi ✅")
    except Exception as e:
        messagebox.showerror("Erreur", str(e))

def export_word():
    try:
        data = historique_text.get("1.0", tk.END).strip()
        if not data:
            messagebox.showwarning("Attention", "Aucune donnée à exporter.")
            return
        filepath = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Fichier Word", "*.docx")])
        if filepath:
            doc = Document()
            doc.add_heading("Rapport des Prédictions", level=1)
            for line in data.split("\n"):
                doc.add_paragraph(line)
            doc.save(filepath)
            messagebox.showinfo("Succès", "Export Word réussi ✅")
    except Exception as e:
        messagebox.showerror("Erreur", str(e))

# -----------------------------
# 6️⃣ - Boutons
# -----------------------------
frame_buttons = tk.Frame(root, bg=BG_COLOR)
frame_buttons.pack(pady=15)

tk.Button(frame_buttons, text=" Prédire", command=predire, font=BUTTON_FONT,
          bg=BTN_COLOR, fg=BTN_TEXT_COLOR, width=15).grid(row=0, column=0, padx=10)

tk.Button(frame_buttons, text="Réinitialiser", command=reset, font=BUTTON_FONT,
          bg="#dc3545", fg="white", width=15).grid(row=0, column=1, padx=10)

frame_export = tk.Frame(root, bg=BG_COLOR)
frame_export.pack(pady=10)

tk.Button(frame_export, text=" Import CSV", command=import_csv, font=BUTTON_FONT,
          bg="#6c757d", fg="white", width=15).grid(row=0, column=0, padx=8)
tk.Button(frame_export, text="Export CSV", command=export_csv, font=BUTTON_FONT,
          bg="#17a2b8", fg="white", width=15).grid(row=0, column=1, padx=8)
tk.Button(frame_export, text="Export PDF", command=export_pdf, font=BUTTON_FONT,
          bg="#28a745", fg="white", width=15).grid(row=0, column=2, padx=8)
tk.Button(frame_export, text=" Export Word", command=export_word, font=BUTTON_FONT,
          bg="#ffc107", fg="black", width=15).grid(row=0, column=3, padx=8)

# -----------------------------
# 7️⃣ - Résultat et Historique
# -----------------------------
label_result = tk.Label(root, text="", font=("Helvetica", 14, "bold"), bg=BG_COLOR)
label_result.pack(pady=10)

tk.Label(root, text=" Historique des prédictions :", font=LABEL_FONT, bg=BG_COLOR).pack(pady=5)

historique_text = scrolledtext.ScrolledText(root, width=80, height=12, wrap=tk.WORD, font=("Courier", 10))
historique_text.pack(padx=15, pady=10)

# -----------------------------
# 8️⃣ - Lancement
# -----------------------------
root.mainloop()
