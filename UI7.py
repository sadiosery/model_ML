# ui6_streamlit_full.py
import streamlit as st
import pandas as pd
from sklearn.preprocessing import OneHotEncoder
from sklearn.ensemble import RandomForestRegressor
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import io

# -----------------------------
# 1️⃣ - Préparer les données
# -----------------------------
data = pd.DataFrame({
    'Jour': ['Lundi','Lundi','Mardi','Mercredi','Jeudi','Vendredi'],
    'Heure': ['Matin','Soir','Matin','Midi','Soir','Matin'],
    'Meteo': ['soleil','pluie','soleil','neige','pluie','soleil'],
    'Promo': ['Oui','Non','Non','Oui','Oui','Non'],
    'Evenement': ['Non','Oui','Non','Non','Oui','Non'],
    'Clients': [120,180,140,100,200,160]
})

encoder = OneHotEncoder(sparse_output=False)
X_cat = encoder.fit_transform(data[['Jour','Heure','Meteo','Promo','Evenement']])
X_cat_df = pd.DataFrame(X_cat, columns=encoder.get_feature_names_out())
X = X_cat_df
y = data['Clients']

model = RandomForestRegressor(n_estimators=100, random_state=42)
model.fit(X, y)

# -----------------------------
# 2️⃣ - Interface Streamlit
# -----------------------------
st.set_page_config(page_title="Prédiction Clients", layout="wide")
st.title("📊 Prédiction du nombre de clients")

# Sélections
jours = ['Lundi','Mardi','Mercredi','Jeudi','Vendredi']
heures = ['Matin','Midi','Soir']
meteo_options = ['soleil','pluie','neige']
promo_options = ['Oui','Non']
evenement_options = ['Oui','Non']

col1, col2, col3 = st.columns(3)

with col1:
    jour = st.selectbox("Jour", jours)
    heure = st.selectbox("Heure", heures)
with col2:
    meteo = st.selectbox("Météo", meteo_options)
    promo = st.selectbox("Promotion", promo_options)
with col3:
    evenement = st.selectbox("Événement", evenement_options)

# Historique dans la session
if "historique" not in st.session_state:
    st.session_state.historique = []

# -----------------------------
# 3️⃣ - Prédiction
# -----------------------------
if st.button("Prédire"):
    try:
        nouveau_jour = pd.DataFrame({
            'Jour': [jour],
            'Heure': [heure],
            'Meteo': [meteo],
            'Promo': [promo],
            'Evenement': [evenement]
        })
        nouveau_jour_cat = encoder.transform(nouveau_jour)
        nouveau_jour_df = pd.DataFrame(nouveau_jour_cat, columns=encoder.get_feature_names_out())
        resultat = model.predict(nouveau_jour_df)
        prediction = int(resultat[0])

        st.success(f"✅ Clients prévus : {prediction}")
        # Ajouter à l'historique
        ligne = f"{jour} - {heure} - {meteo} - Promo:{promo} - Événement:{evenement} ➜ {prediction} clients"
        st.session_state.historique.append(ligne)

    except Exception as e:
        st.error(f"Une erreur est survenue : {e}")

# -----------------------------
# 4️⃣ - Historique
# -----------------------------
st.subheader("📜 Historique des prédictions")
if st.session_state.historique:
    for h in st.session_state.historique:
        st.write(h)
else:
    st.write("Aucune prédiction pour le moment.")

# -----------------------------
# 5️⃣ - Export CSV, PDF, Word
# -----------------------------
# Export CSV
if st.button("Exporter CSV"):
    if st.session_state.historique:
        df_hist = pd.DataFrame(st.session_state.historique, columns=["Historique"])
        csv = df_hist.to_csv(index=False).encode('utf-8')
        st.download_button("Télécharger CSV", data=csv, file_name="historique.csv", mime="text/csv")
    else:
        st.warning("Aucune donnée à exporter.")

# Export PDF
if st.button("Exporter PDF"):
    if st.session_state.historique:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = [Paragraph("Rapport des Prédictions", styles['Title']), Spacer(1, 12)]
        for line in st.session_state.historique:
            story.append(Paragraph(line, styles['Normal']))
        doc.build(story)
        st.download_button("Télécharger PDF", data=buffer.getvalue(), file_name="historique.pdf", mime="application/pdf")
    else:
        st.warning("Aucune donnée à exporter.")

# Export Word
if st.button("Exporter Word"):
    if st.session_state.historique:
        buffer = io.BytesIO()
        docx_doc = Document()
        docx_doc.add_heading("Rapport des Prédictions", level=1)
        for line in st.session_state.historique:
            docx_doc.add_paragraph(line)
        docx_doc.save(buffer)
        st.download_button("Télécharger Word", data=buffer.getvalue(), file_name="historique.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.warning("Aucune donnée à exporter.")

# -----------------------------
# 6️⃣ - Réinitialiser
# -----------------------------
if st.button("Réinitialiser historique"):
    st.session_state.historique = []
    st.success("Historique réinitialisé ✅")
